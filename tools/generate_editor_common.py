#!/usr/bin/env python3
"""Generează pentru-editori/Decizii-structurale-comune.docx.

Memo scurt, doar de citit (nu se completează), cu deciziile structurale care
traversează capitolele — distilat din CLAUDE.md („Decizii deja luate” și
„Ierarhia capitolelor”). Conținutul e hardcodat aici (text scurt și stabil);
dacă deciziile structurale se schimbă, se editează direct această listă.

Artefact generat — .docx-ul nu se editează manual (se editează acest script).

Rulează:  python3 tools/generate_editor_common.py
"""
import os

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
REPO_ROOT = os.path.join(HERE, "..")
OUT_DIR = os.path.join(REPO_ROOT, "pentru-editori")
OUT_PATH = os.path.join(OUT_DIR, "Decizii-structurale-comune.docx")

TITLU = "Decizii structurale comune — toate capitolele"
SUBTITLU = ("Revizia ghidului de indicații radioimagistice aprobat în 2021 · "
            "memo pentru editori")

ACCENT = RGBColor(0x1C, 0x5D, 0x8C)
ACCENT_HEX = "1C5D8C"
GRI = RGBColor(0x55, 0x55, 0x55)
ALB = RGBColor(0xFF, 0xFF, 0xFF)
FONT = "Calibri"

INTRO = (
    "Documentul de față rezumă deciziile de structură aplicate uniform în toate "
    "capitolele reviziei: consolidarea procedurilor intervenționale, regula de "
    "încadrare a unei situații clinice într-un singur capitol, coloana nouă "
    "„Terapeutic” și convențiile de ordonare. Sunt decizii de organizare a "
    "conținutului, nu de conținut clinic."
)
INTRO2 = (
    "Vă rugăm să îl parcurgeți înainte de pachetul dumneavoastră de capitol: "
    "explică de ce anumite situații clinice nu mai apar acolo unde erau în "
    "versiunea din 2021. Nu necesită completare — deciziile care vă privesc "
    "direct sunt formulate ca întrebări în pachetul de capitol."
)

# (titlu, [paragrafe], [bullets])
SECTIUNI = [
    (
        "Ierarhia capitolelor — o situație clinică, un singur capitol",
        ["Ghidul urmează felul în care se pune problema în fața unui caz — "
         "întâi contextul, apoi anatomia. Când o situație clinică poate fi "
         "încadrată în mai multe capitole, se aplică prima regulă care se "
         "potrivește:"],
        [],
    ),
    (
        None,  # continuare a secțiunii de mai sus, lista numerotată
        [],
        [],
    ),
    (
        "Radiologie intervențională — capitol unic pentru proceduri",
        ["Procedurile intervenționale (biopsii, embolizări, drenaje, angiografii "
         "terapeutice) sunt consolidate într-un capitol propriu, organizat pe "
         "aparate și sisteme. Motivul: sunt proceduri terapeutice, nu doar "
         "opțiuni de diagnostic — locul lor nu e întotdeauna potrivit în lista "
         "de investigații a unei situații clinice. Două excepții:"],
        ["ERCP rămâne în capitolul de origine, marcat „Terapeutic = Da”.",
         "Intervenționalul mamar (biopsie percutanată, biopsie ganglionară axilară, "
         "localizare preoperatorie, drenaj de abces) rămâne în capitolul „Sân”, "
         "lângă restul evaluării senologice."],
    ),
    (
        "Coloana nouă „Terapeutic” (Da / Nu)",
        ["Marchează dacă investigația are rol terapeutic, chiar dacă este și "
         "diagnostică. Nu există o a treia valoare „mixt”: embolizarea, ablația "
         "și drenajul sunt „Da”; biopsia și angiografia diagnostică sunt „Nu”."],
        [],
    ),
    (
        "Identificatorul NR.CRT nu are semnificație clinică",
        ["NR.CRT este un simplu identificator de rând, folosit intern. Poate fi "
         "renumerotat integral la finalul reviziei, deci nu este stabil între "
         "versiuni. În corespondență și în comentarii, vă rugăm să vă referiți la "
         "o situație clinică prin „Capitol + Situație clinică”, nu prin număr."],
        [],
    ),
    (
        "O situație clinică = un singur „acasă”",
        ["Fiecare situație clinică apare o singură dată, într-un singur "
         "capitol și subcapitol. Trimiterile încrucișate din date („vezi și "
         "capitolul X”) au fost eliminate: navigarea după organ, peste capitole, "
         "este sarcina aplicației de consultare, nu a textului."],
        [],
    ),
    (
        "Ordinea situațiilor clinice",
        ["În fiecare subcapitol, situațiile clinice sunt ordonate după fluxul "
         "clinic — urgent/acut, apoi cronic, apoi screening și populații speciale "
         "— nu alfabetic. „ALTĂ SITUAȚIE CLINICĂ” rămâne întotdeauna ultima."],
        [],
    ),
    (
        "Ce nu s-a modificat",
        ["Gradele de indicație, dozele și ordinea examenelor în interiorul unei "
         "situații clinice au rămas neatinse. Acolo unde ceva pare discutabil, nu "
         "a fost corectat unilateral, ci este propus ca problemă de decis în "
         "pachetul de capitol."],
        [],
    ),
]

IERARHIE = [
    "Pacient pediatric → capitolul Pediatrie.",
    "Procedură intervențională → capitolul Radiologie intervențională.",
    "Malignitate cunoscută sau suspectată → capitolul Cancer.",
    "Traumatism acut → capitolul Traumatisme.",
    "În rest → capitolul anatomic (organ / aparat).",
]

IERARHIE_NOTA = (
    "Prin urmare, conținutul pediatric, oncologic, traumatic sau intervențional "
    "care era în capitolul dumneavoastră anatomic a fost mutat — nu eliminat. "
    "Senologia face excepție, ca specialitate de sine stătătoare: tot ce ține de "
    "sân, inclusiv cancerul de sân și procedurile intervenționale mamare, rămâne "
    "în capitolul „Sân”."
)

BOARD_INTRO = (
    "Structura actuală — contextul înaintea anatomiei — rămâne neschimbată pentru "
    "această revizie. Semnalăm însă, pentru o discuție separată de board după "
    "încheierea reviewului, câteva fragmentări transversale moștenite:"
)

FRAGMENTARI = [
    "Axul neurologic este împărțit între „Cap › Neuro” și „Coloană vertebrală” — "
    "encefalul, separat de măduvă și coloană.",
    "Sfera ORL este împărțită între „Cap › ORL” și „Gât (părți moi)”.",
    "Patologia vasculară este dispersată: cordul la Cardiovascular, carotidele la "
    "Neuro, teritoriul periferic la Radiologie intervențională.",
    "Patologia endocrină este dispersată: tiroida la Gât, suprarenalele la "
    "Uro-genital, hipofiza la Cap.",
]

BOARD_OUTRO = (
    "Direcția pe care o propunem spre analiză: un capitol „Sistem nervos” "
    "(encefal, măduvă, coloană vertebrală) și un capitol „Cap și gât” (ORL, "
    "tiroidă, mase cervicale). Nimic din acestea nu se decide acum și nu "
    "afectează pachetul de față."
)

INCHEIERE = ("Orice observație sau propunere suplimentară este binevenită și va fi "
             "supusă discuției comune.")


# --- anexă: legenda coloanelor ------------------------------------------

LEGENDA_TITLU = "Anexă — cum se citesc coloanele ghidului"

LEGENDA_INTRO = (
    "Anexa explică sensul coloanelor pe care le veți întâlni cel mai des la "
    "revizuirea capitolului: statutul recomandării, nivelul de dovezi, nivelul "
    "de iradiere și delimitarea dintre cele două coloane de text liber. Este "
    "material de referință — nimic din ea nu se completează."
)

LEG_INDICATIE_INTRO = (
    "Coloana „Indicație” răspunde la o singură întrebare: în situația clinică "
    "descrisă, ce statut are această investigație? Nu este un scor de utilitate "
    "a examenului în general, ci raportat la acel tablou clinic. În ghid apar "
    "șase valori:"
)

LEG_INDICATIE = [
    ("Indicat",
     "Investigația cu cea mai mare șansă de a contribui la diagnostic și la "
     "conduita terapeutică — prima alegere pentru acest tablou clinic."),
    ("Doar în cazuri particulare",
     "Nu este o investigație de rutină. Se ia în calcul doar dacă există un "
     "motiv clinic concret care o justifică; altfel poate fi corect să se amâne "
     "examinarea și să se reevalueze pacientul în timp."),
    ("Doar cu aviz specializat",
     "Investigație complexă, laborioasă și consumatoare de resurse — se "
     "solicită de regulă după discuție cu medicul radiolog sau de medicină "
     "nucleară, ori conform protocolului local."),
    ("Este necesară justificare detaliată",
     "Se efectuează doar dacă medicul solicitant oferă un motiv clinic explicit "
     "și documentat; fără acesta, examinarea nu este justificată."),
    ("Neindicat în primă intenție",
     "Poate deveni utilă ulterior, dar nu ca prim pas — se recomandă epuizarea "
     "investigațiilor mai potrivite înaintea acesteia."),
    ("Neindicat",
     "În acest context clinic, investigația nu este susținută de dovezi."),
]

LEG_GRAD_INTRO = (
    "Coloana „Grad indicație” arată pe ce se sprijină recomandarea — cât de "
    "solide sunt dovezile din spatele ei. Este independentă de coloana "
    "„Indicație”: o investigație poate fi „Indicat” cu grad C (consens al "
    "specialiștilor, în lipsa unor studii ample) sau „Neindicat” cu grad A."
)

LEG_GRAD = [
    ("A", "Nivel de dovezi ridicat — studii sau meta-analize solide."),
    ("B", "Nivel de dovezi moderat."),
    ("C", "Recomandare bazată pe consens sau pe dovezi limitate."),
    ("?", "Nivel de dovezi neprecizat în sursa originală."),
    ("(gol)",
     "Rând-punte, fără examen concret (Tip Z) — de exemplu „ALTĂ SITUAȚIE "
     "CLINICĂ”. Nu are grad și nici doză."),
]

LEG_DOZA_INTRO = (
    "Iradierea nu este exprimată în milisievert, ci pe o scală relativă de "
    "nivele, de la 0 la 4, aceeași pentru tot ghidul. „Doza 4” înseamnă nivelul "
    "cel mai ridicat de iradiere, nu 4 mSv. Fiecare examen are o pereche "
    "„Doza Min – Doza Max”, pentru că iradierea reală depinde de protocol, de "
    "aparat și de regiunea explorată: minimul este cazul favorabil, iar "
    "„Doza Max” este nivelul maxim de iradiere — valoarea de luat în calcul "
    "când se cântărește justificarea examenului. Când cele două coincid, "
    "examenul are o iradiere previzibilă."
)

LEG_DOZA = [
    ("0", "Fără iradiere — ecografie, IRM (fără radiații ionizante)."),
    ("1", "Foarte redusă — echivalent cu câteva zile până la câteva săptămâni "
          "de fond natural de radiație."),
    ("2", "Redusă — echivalent cu câteva luni până la un an de fond natural."),
    ("3", "Medie — echivalent cu aproximativ 1,5–2 ani de fond natural."),
    ("4", "Ridicată — peste aproximativ 2 ani de fond natural."),
]

LEG_TEXT_INTRO = (
    "Ghidul are două coloane de text liber, cu roluri diferite. Distincția "
    "contează la revizuire: ce scrieți într-una nu se caută și nu se afișează "
    "la fel ca în cealaltă."
)

LEG_TEXT = [
    ("Comentarii",
     "Conținut clinic: precizări despre situație, criterii de alegere între "
     "examene, cadența de urmărire, diagnostic diferențial, condiții tehnice "
     "care schimbă indicația. Tot ce ajută medicul să decidă."),
    ("Alte informații",
     "Coduri interne și trimiteri bibliografice: codurile de procedură din "
     "radiologia intervențională (de forma „RI - PBx”) și sursele invocate "
     "(ACR Appropriateness Criteria, ghiduri ESMO, RCR și altele)."),
]

LEG_TEXT_NOTA = (
    "O observație care evită o confuzie frecventă: informația clinică valabilă "
    "pentru toată situația este repetată intenționat, în „Comentarii”, pe "
    "fiecare investigație a acelei situații. Nu este o scăpare de redactare — "
    "rândurile se consultă separat, câte unul, iar fiecare trebuie să fie "
    "inteligibil de sine stătător. Doar nota strict specifică unei modalități "
    "(ce arată sau cum se face acel examen) rămâne pe rândul ei."
)


# --- helpers docx --------------------------------------------------------

def set_base_style(doc):
    """Font și spațiere de bază, uniforme pentru tot documentul."""
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = normal.paragraph_format
    pf.space_after = Pt(8)
    pf.line_spacing = 1.15

    for name, size, before, after in (
        ("Heading 1", 20, 0, 6),
        ("Heading 2", 13, 14, 4),
        ("List Bullet", 11, 0, 4),
        ("List Number", 11, 0, 4),
    ):
        st = doc.styles[name]
        st.font.name = FONT
        st.font.size = Pt(size)
        st.font.color.rgb = ACCENT if name.startswith("Heading") else None
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        if name.startswith("Heading"):
            st.font.bold = True


def set_page(doc):
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)  # A4
    section.top_margin = section.bottom_margin = Cm(2.0)
    section.left_margin = section.right_margin = Cm(2.2)


def add_field(paragraph, instr):
    """Inserează un câmp Word (ex. PAGE / NUMPAGES) într-un paragraf."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr_el, end):
        run._r.append(el)
    return run


def add_footer(doc):
    """Subsol: titlul scurt la stânga, „Pagina X din Y” la dreapta."""
    section = doc.sections[0]
    p = section.footer.paragraphs[0]
    # Stilul „Footer” aduce tabulatori moșteniți (centru 4680, dreapta 9360
    # twips, pentru pagină Letter); îi anulăm ca tabul să ajungă la marginea A4.
    for pos in (Pt(234), Pt(468)):  # 4680 / 9360 twips
        p.paragraph_format.tab_stops.add_tab_stop(pos, WD_TAB_ALIGNMENT.CLEAR)
    p.paragraph_format.tab_stops.add_tab_stop(
        section.page_width - section.left_margin - section.right_margin,
        WD_TAB_ALIGNMENT.RIGHT,
    )
    run = p.add_run("Decizii structurale comune\t")
    run.font.size = Pt(9)
    run.font.color.rgb = GRI
    pag = p.add_run("Pagina ")
    pag.font.size = Pt(9)
    pag.font.color.rgb = GRI
    for instr, sep in (("PAGE", " din "), ("NUMPAGES", "")):
        f = add_field(p, instr)
        f.font.size = Pt(9)
        f.font.color.rgb = GRI
        if sep:
            r = p.add_run(sep)
            r.font.size = Pt(9)
            r.font.color.rgb = GRI


def add_rule(doc):
    """Linie orizontală subțire, în culoarea de accent."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1C5D8C")
    pbdr.append(bottom)
    p._p.get_or_add_pPr().append(pbdr)
    return p


def add_heading(doc, text, level=2, number=None):
    label = f"{number}. {text}" if number else text
    return doc.add_heading(label, level=level)


def add_note(doc, text):
    """Paragraf de notă — gri, italic, ușor indentat."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = GRI
    return p


def set_cell_text(cell, text, bold=False, size=10):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    cell.text = ""
    par = cell.paragraphs[0]
    # `cell.text = ""` lasă un run gol; îl scoatem ca par.runs[0] să fie textul.
    for r in list(par.runs):
        r._r.getparent().remove(r._r)
    par.paragraph_format.space_after = Pt(2)
    par.paragraph_format.line_spacing = 1.05
    run = par.add_run(text or "")
    run.bold = bold
    run.font.size = Pt(size)
    return par


def set_col_widths(table, widths):
    """Fixează lățimile coloanelor (în Pt). Lățimea trebuie scrisă în trei
    locuri: pe fiecare celulă (python-docx ignoră columns[].width), în
    <w:tblGrid> și ca lățime totală de tabel."""
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = Pt(w)
    twips = [int(round(w * 20)) for w in widths]
    grid = table._tbl.find(qn("w:tblGrid"))
    for col, tw in zip(grid.findall(qn("w:gridCol")), twips):
        col.set(qn("w:w"), str(tw))
    tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
    if tbl_w is not None:
        tbl_w.set(qn("w:type"), "dxa")
        tbl_w.set(qn("w:w"), str(sum(twips)))


def shade_cell(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def legend_table(doc, headers, rows):
    """Tabel de legendă pe două coloane, antet alb pe fond de accent.

    Lățimile însumează 470 Pt = 16,6 cm, exact oglinda utilă a paginii A4
    portret cu marginile documentului (21,0 − 2×2,2 cm).
    """
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    head = table.rows[0]
    for cell, text in zip(head.cells, headers):
        par = set_cell_text(cell, text, bold=True)
        par.runs[0].font.color.rgb = ALB
        shade_cell(cell, ACCENT_HEX)
    tr_pr = head._tr.get_or_add_trPr()
    rep = OxmlElement("w:tblHeader")
    rep.set(qn("w:val"), "true")
    tr_pr.append(rep)
    for eticheta, descriere in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], eticheta, bold=True)
        set_cell_text(cells[1], descriere)
    set_col_widths(table, [125, 345])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def set_docprops(doc):
    core = doc.core_properties
    core.title = TITLU
    core.subject = "Revizia ghidului de indicații radioimagistice"
    core.category = "Memo pentru editori"
    core.comments = "Artefact generat de tools/generate_editor_common.py"


# --- document ------------------------------------------------------------

def build(doc):
    set_base_style(doc)
    set_page(doc)
    set_docprops(doc)
    add_footer(doc)

    doc.add_heading(TITLU, level=1)
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(2)
    run = sub.add_run(SUBTITLU)
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRI
    run.bold = True
    add_rule(doc)

    doc.add_paragraph(INTRO)
    doc.add_paragraph(INTRO2)

    numar = 0
    for titlu, paragrafe, bullets in SECTIUNI:
        if titlu is None:
            for item in IERARHIE:
                doc.add_paragraph(item, style="List Number")
            add_note(doc, IERARHIE_NOTA)
            continue
        numar += 1
        add_heading(doc, titlu, level=2, number=numar)
        for text in paragrafe:
            doc.add_paragraph(text)
        for item in bullets:
            doc.add_paragraph(item, style="List Bullet")

    numar += 1
    add_heading(doc, "Întrebare de board, după încheierea reviewului: structura de capitole",
                level=2, number=numar)
    doc.add_paragraph(BOARD_INTRO)
    for frag in FRAGMENTARI:
        doc.add_paragraph(frag, style="List Bullet")
    doc.add_paragraph(BOARD_OUTRO)

    add_rule(doc)
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = closing.add_run(INCHEIERE)
    run.italic = True
    run.font.color.rgb = GRI

    build_legenda(doc)


def build_legenda(doc):
    """Anexa finală, pe pagină proprie: sensul coloanelor ghidului."""
    doc.add_page_break()
    doc.add_heading(LEGENDA_TITLU, level=1)
    add_rule(doc)
    doc.add_paragraph(LEGENDA_INTRO)

    doc.add_heading("Indicație — statutul recomandării", level=2)
    doc.add_paragraph(LEG_INDICATIE_INTRO)
    legend_table(doc, ["Valoare", "Ce înseamnă"], LEG_INDICATIE)

    doc.add_heading("Grad indicație — nivelul de dovezi", level=2)
    doc.add_paragraph(LEG_GRAD_INTRO)
    legend_table(doc, ["Grad", "Ce înseamnă"], LEG_GRAD)

    doc.add_heading("Doza minimă și doza maximă — nivelul de iradiere", level=2)
    doc.add_paragraph(LEG_DOZA_INTRO)
    legend_table(doc, ["Nivel", "Ce înseamnă"], LEG_DOZA)

    doc.add_heading("„Comentarii” față de „Alte informații”", level=2)
    doc.add_paragraph(LEG_TEXT_INTRO)
    legend_table(doc, ["Coloană", "Ce conține"], LEG_TEXT)
    add_note(doc, LEG_TEXT_NOTA)


def main():
    doc = Document()
    build(doc)
    os.makedirs(OUT_DIR, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"✓ {OUT_PATH}")


if __name__ == "__main__":
    main()
