#!/usr/bin/env python3
"""Generează pachetele .docx de trimis editorilor, câte unul per capitol.

Fiecare pachet conține: modificări deja aplicate (informativ), discuții
deschise, propuneri noi și duplicate de rezolvat — preluate dintr-un
fișier YAML de conținut distilat, unul per capitol, în
tools/editor_content/<NN-capitol>.yaml. Conținutul curent al capitolului NU
e reprodus în docx (vine din GHID.csv, copiat separat înainte de trimitere).

Rulează:
    python3 tools/generate_editor_docx.py pediatrie cardiovascular
    python3 tools/generate_editor_docx.py --all

Scrie în pentru-editori/<NN-capitol>.docx — artefact generat,
nu se editează manual (se editează YAML-ul sursă + acest script).
"""
import argparse
import os
import sys

import yaml
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
REPO_ROOT = os.path.join(HERE, "..")
CONTENT_DIR = os.path.join(HERE, "editor_content")
OUTPUT_ROOT = os.path.join(REPO_ROOT, "pentru-editori")

# (slug, Capitol exact ca în GHID.csv, nume fișier)
CHAPTERS = [
    ("01-pediatrie", "Pediatrie", "Pediatrie"),
    ("02-traumatisme", "Traumatisme", "Traumatisme"),
    ("03-cancer", "Cancer", "Cancer"),
    ("04-aparat-cardiovascular", "Aparat cardiovascular", "Aparat cardiovascular"),
    ("05-torace", "Torace", "Torace"),
    ("06-aparat-digestiv", "Aparat digestiv", "Aparat digestiv"),
    ("07-aparat-uro-genital-si-glande-suprarenale",
     "Aparat uro-genital și glande suprarenale",
     "Aparat uro-genital și glande suprarenale"),
    ("08-obstetrica-si-ginecologie", "Obstetrică și ginecologie", "Obstetrică și ginecologie"),
    ("09-san", "Sân", "Sân"),
    ("10-cap", "Cap", "Cap"),
    ("11-gat-parti-moi", "Gât (părți moi)", "Gât (părți moi)"),
    ("12-coloana-vertebrala", "Coloană vertebrală", "Coloană vertebrală"),
    ("13-aparat-locomotor", "Aparat locomotor", "Aparat locomotor"),
    ("14-radiologie-interventionala", "Radiologie intervențională", "Radiologie intervențională"),
]
SLUG_TO_CHAPTER = {slug: (cap, fname) for slug, cap, fname in CHAPTERS}
CONTENT_KEY_TO_SLUG = {
    "pediatrie": "01-pediatrie",
    "traumatisme": "02-traumatisme",
    "cancer": "03-cancer",
    "cardiovascular": "04-aparat-cardiovascular",
    "torace": "05-torace",
    "digestiv": "06-aparat-digestiv",
    "urogenital": "07-aparat-uro-genital-si-glande-suprarenale",
    "obstetrica": "08-obstetrica-si-ginecologie",
    "san": "09-san",
    "cap": "10-cap",
    "gat": "11-gat-parti-moi",
    "coloana": "12-coloana-vertebrala",
    "locomotor": "13-aparat-locomotor",
    "ri": "14-radiologie-interventionala",
}

ACCENT = RGBColor(0x1C, 0x5D, 0x8C)
ACCENT_HEX = "1C5D8C"
GRI = RGBColor(0x55, 0x55, 0x55)
ALB = RGBColor(0xFF, 0xFF, 0xFF)
FONT = "Calibri"

# Lățimi de coloană (Pt), calculate pentru A4 landscape cu margini de 1.5 cm
# (29,7 - 2 x 1,5 = 26,7 cm = 757 pt utili).
W_TABEL_3 = [425, 150, 180]
W_TABEL_8 = [32, 118, 95, 38, 40, 170, 105, 158]


def load_content(slug):
    path = os.path.join(CONTENT_DIR, f"{slug}.yaml")
    if not os.path.exists(path):
        return {}
    with open(path, encoding="utf-8") as f:
        return yaml.safe_load(f) or {}


# --- helpers docx --------------------------------------------------------

def set_base_style(doc):
    """Font și spațiere de bază, uniforme pentru tot documentul."""
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before, after in (
        ("Heading 1", 20, 0, 6),
        ("Heading 2", 13, 16, 6),
        ("Heading 3", 11.5, 12, 4),
        ("List Bullet", 11, 0, 4),
    ):
        st = doc.styles[name]
        st.font.name = FONT
        st.font.size = Pt(size)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        if name.startswith("Heading"):
            st.font.bold = True
            st.font.color.rgb = ACCENT
            st.paragraph_format.keep_with_next = True


def set_page(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = Cm(29.7), Cm(21.0)  # A4 landscape
    section.top_margin = section.bottom_margin = Cm(1.5)
    section.left_margin = section.right_margin = Cm(1.5)


def set_docprops(doc, chapter_name):
    core = doc.core_properties
    core.title = f"Pachet de review — {chapter_name}"
    core.subject = "Revizia ghidului de indicații radioimagistice"
    core.category = "Pachet de capitol pentru editori"
    core.comments = "Artefact generat de tools/generate_editor_docx.py"


def add_field(paragraph, instr):
    """Inserează un câmp Word (ex. PAGE / NUMPAGES) într-un paragraf."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr_el, end):
        run._r.append(el)
    return run


def add_footer(doc, chapter_name):
    """Subsol: numele capitolului la stânga, „Pagina X din Y” la dreapta."""
    section = doc.sections[0]
    p = section.footer.paragraphs[0]
    # Stilul „Footer” aduce tabulatori moșteniți (centru 4680, dreapta 9360
    # twips, pentru pagină Letter); îi anulăm ca tabul să ajungă la marginea
    # din dreapta a paginii A4 landscape.
    for pos in (Pt(234), Pt(468)):
        p.paragraph_format.tab_stops.add_tab_stop(pos, WD_TAB_ALIGNMENT.CLEAR)
    p.paragraph_format.tab_stops.add_tab_stop(
        section.page_width - section.left_margin - section.right_margin,
        WD_TAB_ALIGNMENT.RIGHT,
    )

    def small(text):
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = GRI
        return run

    small(f"{chapter_name} — pachet de review\t")
    small("Pagina ")
    for instr, sep in (("PAGE", " din "), ("NUMPAGES", "")):
        f = add_field(p, instr)
        f.font.size = Pt(9)
        f.font.color.rgb = GRI
        if sep:
            small(sep)


def add_rule(doc):
    """Linie orizontală subțire, în culoarea de accent."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), ACCENT_HEX)
    pbdr.append(bottom)
    p._p.get_or_add_pPr().append(pbdr)
    return p


def add_note(doc, text):
    """Paragraf de notă — gri, italic."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = GRI
    return p


def set_cell_text(cell, text, bold=False, italic=False, size=9.5):
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p = cell.paragraphs[0]
    # `cell.text = ""` lasă un run gol; îl scoatem ca p.runs[0] să fie textul.
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text or "")
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p


def set_cell_paragraph(cell, text, size=9.5):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def shade_cell(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def new_table(doc, headers, widths):
    """Tabel cu grilă simplă, antet colorat care se repetă la schimbarea paginii."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    for cell, text in zip(header.cells, headers):
        p = set_cell_text(cell, text, bold=True)
        p.runs[0].font.color.rgb = ALB
        shade_cell(cell, ACCENT_HEX)
    repeat_header_row(header)
    set_col_widths(table, widths)
    return table


def repeat_header_row(row):
    """Repetă rândul de antet pe fiecare pagină (fără API în python-docx)."""
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_col_widths(table, widths):
    """Fixează lățimile coloanelor (în Pt). Lățimea trebuie scrisă în trei
    locuri: pe fiecare celulă (python-docx ignoră columns[].width), în
    <w:tblGrid> (altfel Word/LibreOffice folosesc grila implicită, egală) și
    ca lățime totală de tabel."""
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = Pt(w)

    twips = [int(round(w * 20)) for w in widths]
    grid = table._tbl.find(qn("w:tblGrid"))
    for col, tw in zip(grid.findall(qn("w:gridCol")), twips):
        col.set(qn("w:w"), str(tw))
    tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
    if tbl_w is not None:
        tbl_w.set(qn("w:type"), "dxa")
        tbl_w.set(qn("w:w"), str(sum(twips)))


# --- secțiuni --------------------------------------------------------

def add_intro(doc, chapter_name):
    doc.add_heading(f"{chapter_name} — pachet de review", level=1)
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(2)
    run = sub.add_run(
        "Revizia ghidului de indicații radioimagistice aprobat în 2021 · "
        "document de lucru pentru editorul de capitol"
    )
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRI
    add_rule(doc)

    doc.add_paragraph(
        "Documentul cuprinde, pentru capitolul dumneavoastră: modificările deja "
        "aplicate (informativ), problemele rămase de decis, propunerile de "
        "conținut nou și duplicatele de rezolvat. Conținutul actual al "
        "capitolului nu este reprodus aici — vine din fișierele GHID.csv și "
        "GHID.html, atașate."
    )
    doc.add_paragraph(
        "Vă rugăm să completați coloanele „Decizie” (Aprobat / Respins / Amânat) "
        "și „Comentariu” direct în tabele, fără să le modificați structura. "
        "Acolo unde o problemă are variante enumerate, este suficient să "
        "indicați varianta aleasă."
    )
    add_note(
        doc,
        "Atașate: GHID.csv și GHID.html — conținutul integral al ghidului — și "
        "memoul „Decizii structurale comune”, de citit înaintea acestui pachet.",
    )


def add_fyi(doc, items):
    add_heading(doc, "1. Modificări deja aplicate (informativ)", level=2)
    if not items:
        add_note(doc, "Nicio modificare specifică acestui capitol.")
        return
    doc.add_paragraph(
        "Nu necesită decizie; sunt listate pentru ca revizuirea capitolului să "
        "pornească de la starea curentă."
    )
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_open_questions(doc, items):
    add_heading(doc, "2. Probleme de decis", level=2)
    if not items:
        add_note(doc, "Nicio problemă deschisă în acest capitol.")
        return
    table = new_table(doc, ["Problemă", "Decizie", "Comentariu"], W_TABEL_3)
    for item in items:
        cells = table.add_row().cells
        p = set_cell_text(cells[0], item["titlu"], bold=True)
        p.paragraph_format.space_after = Pt(3)
        if item.get("descriere"):
            set_cell_paragraph(cells[0], item["descriere"])
        for opt in item.get("optiuni", []):
            set_cell_paragraph(cells[0], f"• {opt}")
        set_cell_text(cells[1], "")
        set_cell_text(cells[2], "")
    set_col_widths(table, W_TABEL_3)


def add_draft_rows(doc, groups):
    add_heading(doc, "3. Propuneri de conținut nou", level=2)
    if not groups:
        add_note(doc, "Nicio propunere nouă în acest capitol.")
        return
    doc.add_paragraph(
        "Gradele și dozele sunt orientative — necesită validare. Fiecare rând, "
        "separat, se aprobă, se respinge sau se amână."
    )
    for group in groups:
        add_heading(doc, f"{group['subcapitol']} — {group['situatie']}", level=3)
        if group.get("referinta"):
            add_note(doc, group["referinta"])
        headers = ["Tip", "Examen", "Indicație", "Grad", "Doză", "Comentarii",
                   "Decizie", "Comentariu editor"]
        table = new_table(doc, headers, W_TABEL_8)
        for row in group["indicatii"]:
            cells = table.add_row().cells
            for cell, key in zip(cells, ("tip", "examen", "indicatie", "grad",
                                         "doza", "comentarii")):
                set_cell_text(cell, row.get(key, ""))
            set_cell_text(cells[6], "")
            set_cell_text(cells[7], "")
        set_col_widths(table, W_TABEL_8)
        doc.add_paragraph()


def add_duplicates(doc, items):
    add_heading(doc, "4. Duplicate de rezolvat", level=2)
    if not items:
        add_note(doc, "Niciun duplicat deschis pentru acest capitol.")
        return
    doc.add_paragraph(
        "Situații apropiate ca sens, apărute de două ori. De decis dacă se "
        "comasează, se reformulează sau rămân distincte."
    )
    table = new_table(doc, ["Descriere", "Decizie", "Comentariu"], W_TABEL_3)
    for item in items:
        cells = table.add_row().cells
        set_cell_text(cells[0], item)
        set_cell_text(cells[1], "")
        set_cell_text(cells[2], "")
    set_col_widths(table, W_TABEL_3)


def generate_chapter(slug):
    _, display_name = SLUG_TO_CHAPTER[slug]
    content = load_content(slug)

    doc = Document()
    set_base_style(doc)
    set_page(doc)
    set_docprops(doc, display_name)
    add_footer(doc, display_name)
    add_intro(doc, display_name)
    add_fyi(doc, content.get("fyi", []))
    add_open_questions(doc, content.get("discutii", []))
    add_draft_rows(doc, content.get("propuneri", []))
    add_duplicates(doc, content.get("duplicate", []))

    os.makedirs(OUTPUT_ROOT, exist_ok=True)
    out_path = os.path.join(OUTPUT_ROOT, f"{slug}.docx")
    doc.save(out_path)
    return out_path


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("chapters", nargs="*", help="chei de conținut (ex. pediatrie cardiovascular)")
    parser.add_argument("--all", action="store_true", help="generează toate cele 14 capitole")
    args = parser.parse_args()

    if args.all:
        keys = list(CONTENT_KEY_TO_SLUG.keys())
    elif args.chapters:
        keys = args.chapters
    else:
        parser.error("specifică cel puțin un capitol sau --all")

    for key in keys:
        if key not in CONTENT_KEY_TO_SLUG:
            sys.exit(f"Capitol necunoscut: {key} (opțiuni: {sorted(CONTENT_KEY_TO_SLUG)})")
        slug = CONTENT_KEY_TO_SLUG[key]
        out_path = generate_chapter(slug)
        print(f"✓ {out_path}")


if __name__ == "__main__":
    main()
